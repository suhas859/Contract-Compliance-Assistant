from docx import Document

def parse_docx(file_path: str) -> str:
    """
    Extracts text from a .docx file, including both paragraphs AND
    tables.
    doc.paragraphs alone misses table content entirely a Word table's
    cells live in a separate doc.tables structure. This matters
    concretely here: contracts/invoices with line-item tables.

    """

    doc = Document(file_path)

    parts = []
 
    # Walk paragraphs and tables in document order 
    for element in doc.element.body:
        if element.tag.endswith("}p"):
            for para in doc.paragraphs:
                if para._element is element:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif element.tag.endswith("}tbl"):
            for table in doc.tables:
                if table._element is element:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))
                    break
    return "\n".join(parts)
